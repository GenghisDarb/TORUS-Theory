import sys
import json
from docx import Document
from datetime import datetime
import requests
import os

# Usage: python append_ledger_and_badge.py result.json
json_path = sys.argv[1]
with open(json_path) as f:
    result = json.load(f)

# Append to ledger
docx_path = "TORUS Validation Ledger.docx"
if os.path.exists(docx_path):
    doc = Document(docx_path)
else:
    doc = Document()
    table = doc.add_table(rows=1, cols=len(result))
    hdr_cells = table.rows[0].cells
    for i, k in enumerate(result):
        hdr_cells[i].text = k
    doc.add_table(table)

# Find or create table
if not doc.tables:
    table = doc.add_table(rows=1, cols=len(result))
    hdr_cells = table.rows[0].cells
    for i, k in enumerate(result):
        hdr_cells[i].text = k
else:
    table = doc.tables[0]
row = table.add_row().cells
for i, k in enumerate(result):
    row[i].text = str(result[k])
doc.save(docx_path)

# Update badge
badge_path = "docs/evidence-badge.svg"
label = "TORUS Evidence"
value = str(result.get("sigma", "?"))
color = "blue"
badge_url = f"https://img.shields.io/badge/{label}-{value}-{color}.svg"
badge_url = badge_url.replace(' ', '%20')
r = requests.get(badge_url)
if r.ok:
    os.makedirs("docs", exist_ok=True)
    with open(badge_path, "wb") as f:
        f.write(r.content)
print("Ledger and badge updated.")
